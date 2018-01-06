# coding: utf8
from docx import Document
import io


document = Document('demo.docx')

file = io.open("unnecesary_qualificators.txt", mode="r", encoding="utf-8")
unnecesary_qualificators = set(line.strip() for line in file)
result_unnecesary_qualificators = []

file = io.open("weak_words.txt", mode="r", encoding="utf-8")
weak_words = set(line.strip() for line in file)
result_weak_word = []

file = io.open("redundancy.txt", mode="r", encoding="utf-8")
redundancy = set(line.strip() for line in file)
result_redundancy = []

file = io.open("exagerations.txt", mode="r", encoding="utf-8")
exagerations = set(line.strip() for line in file)
result_exagerations = []


def get_text(filename):
    doc = Document(filename)

    for para in doc.paragraphs:
        for word in unnecesary_qualificators:
            if word in para.text:
                print('[X]: WORD "', word, '" PRESENT in "', para.text, '", as UNNECESARY QUALIFICATOR')
                result_unnecesary_qualificators.append(word)

        for word in weak_words:
            if word in para.text:
                print('[X]: WORD "', word, '" PRESENT in "', para.text, '", as WEAK WORD')
                result_weak_word.append(word)

        for word in redundancy:
            if word in para.text:
                print('[X]: WORD "', word, '" PRESENT in "', para.text, '", as REDUNDANCY')
                result_redundancy.append(word)

        for word in exagerations:
            if word in para.text:
                print('[X]: WORD "', word, '" PRESENT in "', para.text, '", as EXAGERATIONS')
                result_exagerations.append(word)

    print('UNNECESARY QUALIFICATORS: ', result_unnecesary_qualificators)
    print('WEAK WORDS: ', result_weak_word)
    print('REDUNDANCY: ', result_redundancy)
    print('EXAGERATIONS: ', result_exagerations)


get_text('demo.docx')

